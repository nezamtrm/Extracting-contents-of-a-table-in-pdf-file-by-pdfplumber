
# Import the 'Document' class from the 'docx' module
from docx import Document

# Import the 'pdfplumber' module
import pdfplumber

# Import the 'pandas' module and rename it to 'pd'
import pandas as pd

# Define a dictionary containing settings for table extraction from pdf using pdfplumber
table_settings = {"vertical_strategy": 'text', 'horizontal_strategy': 'text'}

# Open the pdf file using pdfplumber and assign the resulting object to variable 'pdf'
pdf = pdfplumber.open('0_deep_survey.pdf')

# Extract table data from page 17 of the pdf and assign it to variable 'table'
table = pdf.pages[17].extract_table(table_settings)

# Create a pandas DataFrame using the table data and assign it to variable 'df_main'
# Ignore the first row as it is the table header, and set the column names using the first row of the table data
df_main = pd.DataFrame(table[1::], columns= table[0])

# Create a new Word document using the 'Document' class from the 'docx' module and assign it to variable 'doc'
doc = Document()

# Add a table to the end of the Word document using the 'add_table()' method of the 'Document' class,
# and assign the resulting object to variable 't'
# Add an extra row to the table so that we can add the header row separately
t = doc.add_table(df_main.shape[0] + 1, df_main.shape[1])

# Add the header row to the table
for j in range(df_main.shape[-1]):
    # Set the text of the cell at row 0, column j to the jth column name of the DataFrame
    t.cell(0, j).text = df_main.columns[j]

# Add the remaining data of the DataFrame to the table
for i in range(df_main.shape[0]):
    for j in range(df_main.shape[-1]):
        # Set the text of the cell at row i+1, column j to the value of the ith row and jth column of the DataFrame
        t.cell(i + 1, j).text = str(df_main.values[i, j])

# Save the Word document to a file named 'wordfile.docx'
doc.save('wordfile.docx')
